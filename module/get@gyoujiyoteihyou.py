import docx

document = docx.Document(r"module\pdf\2023_yotei_zenki.docx")

for paragraph in document.paragraphs:
    print(paragraph.text)

with open(r"module\pdf\output.txt", "w", encoding="utf-8")as f:
    for table1 in document.tables:
        for row1 in table1.rows:
            for cell1 in row1.cells:
                if len(cell1.tables)==0:
                    #print(cell1.text)
                    pass
                else:
                    for table2 in cell1.tables:
                        for row_num, row2 in enumerate(table2.rows):
                            if row_num==0:
                                f.write(row2.cells[0].text + "\n")
                                #print(row2.cells[0].text)
                            elif row_num==31:
                                try:
                                    f.write(row2.cells[31].text + "\n")

                                except IndexError:
                                    pass
                                    #print(len(row2.cells))
                            else:
                                for cell2 in row2.cells:
                                    f.write(repr(cell2.text)[1:len(repr(cell2.text))-1] + "\n")

print(len(document.tables))