import docx #docxファイルの操作
from pdf2docx.main import parse #pdf ->docxファイルに変換

import scrape   #スクレイピングをするモジュール

import time #時間
import os   #ファイル操作
import re
with open(r"module\pdf\access_log.txt", "r")as f:
    before_access=f.readlines()[0]

before_access=[before_access[:4], before_access[4:6], before_access[6:]]
now_time=[time.strftime("%Y"), time.strftime("%m"), time.strftime("%d")]

before_access = [int(i) for i in before_access]
now_time      = [int(i) for i in now_time]

print(before_access, now_time)
if before_access[0]<now_time[0] or before_access[1]<now_time[1] or before_access[2]<now_time[2]:
    print("other day")
    #scrape.scr("https://www.ktc.ac.jp/dept/kyomu/schedule/")    #学校の行事予定のページのスクレイピング
    pdf_folder="module/pdf"
    with open(pdf_folder+"/access_log.txt", "w")as f:   #最終アクセス日の更新
        f.write(time.strftime("%Y%m%d"))

    nenndo=str(int(time.strftime("%Y")) -1) if time.strftime("%m")=="01" or time.strftime("%m")=="02" or time.strftime("%m")=="03" else time.strftime("%Y")
    parse("module/pdf/" + nenndo + "_yotei_zenki.pdf" , "module/pdf/" + nenndo + "_yotei_zenki.docx")   #pdf2docx
    parse("module/pdf/" + nenndo + "_yotei_kouki.pdf" , "module/pdf/" + nenndo + "_yotei_kouki.docx")

    with open("module/pdf/yotei.txt", "w")as f: #空白にする
        f.write("")

        #要素をテキストファイルに書き出す
    for i in ["zenki", "kouki"]:
        document = docx.Document("module/pdf/" + nenndo + "_yotei_"+i +".docx")
        with open("module/pdf/yotei.txt", "a", encoding="UTF-8")as f:
            for table1 in document.tables:
                for row1 in table1.rows:
                    for cell1 in row1.cells:
                        if len(cell1.tables)==0:
                            #print(cell1.text)
                            pass
                        else:
                            for table2 in cell1.tables:
                                for row_num, row2 in enumerate(table2.rows):
                                    if row_num==0:
                                        f.write(row2.cells[0].text + "\n")
                                        #print(row2.cells[0].text)
                                    elif row_num==31:
                                        try:
                                            f.write(row2.cells[31].text + "\n")

                                        except IndexError:
                                            pass
                                            #print(len(row2.cells))
                                    else:
                                        for cell2 in row2.cells:
                                            
                                            try:
                                                f.write(repr(cell2.text)[1:len(repr(cell2.text))-1] + "\n")
                                                #print(cell2.text, file=f)

                                            except UnicodeEncodeError:
                                                pass


        

    with open(r"C:\Users\user\Documents\3I_PBL_C.github.io\module\pdf\yotei.txt", "r", encoding="utf-8")as f:
        get_data=f.readlines()


    maximum_day=[31,28,31,30,31,30,31,31,30,31,30,31]
    text=[]
    index=0
    month=""
    day=""
    add_text=""

    overflaw=[]

    while index< len(get_data):
        if re.fullmatch(" *[0-9]+ *月 *\n", get_data[index]):
            month=get_data[index].replace(" ", "").replace("月", "")
            index+=1

            continue

        elif re.fullmatch(" *[0-9]+ *\n", get_data[index]):
            day=get_data[index].replace(" ", "")

            add_text=month + "月"+ day + "日" + get_data[index+1] + " "  + get_data[index+2]
            index+=3

            """
            try:
                while int(day)!= maximum_day[int(month)-1] and not re.fullmatch(" *[0-9] *", get_data[index]):
                    add_text+=("\n" + get_data[index])
                    index+=1

                if int(day)== maximum_day[int(month)-1]:
                    while not re.fullmatch(" *[0-9]+ *月 *\n", get_data[index]):
                        overflaw.append(get_data[index])
                        index+=1

                    for _ in range(len(overflaw)%3):
                        add_text+=("\n" + overflaw[0].pop)

                text.append(add_text)
                for i in range(0, len(overflaw)/3):
                    text.append((month + "月 " + overflaw[i] + "\n") if overflaw[i] != "" else "")

            except IndexError:
                continue
            """
            if add_text!="":
                text.append((add_text.replace("\n", "") + "\n") if
                (re.fullmatch("[0-9]+月[0-9]+日 \(.\)  ?.?.+", add_text.replace("\n", ""))
                and not re.fullmatch("[0-9]+月[0-9]+日 \(.\) ( ? ?.+\[補講日?・予備日\]| .)", add_text.replace("\n", ""))
                ) else "")
        else:
            #print(get_data[index])
            index+=1


    text_day=[]
    text_after=[]
    for i in text:
        if (i not in text_after) and i!="":
            text_after.append(i)
            try:
                text_day.append(int( ("2" if int(i.split("日")[0].split("月")[0])>3 else "3") + i.split("日")[0].split("月")[0].zfill(2) + i.split("日")[0].split("月")[1].zfill(2)))
            except ValueError:
                print(i, "a", i.split("日")[0])

    text_day_sorted=[]
    text_after_sorted=[]
    for i in range(len(text_day)):
        min_index= text_day.index(min(text_day))
        text_day_sorted.append(text_day.pop(min_index))
        text_after_sorted.append(text_after.pop(min_index))


    with open(r"C:\Users\user\Documents\3I_PBL_C.github.io\module\pdf\output.txt", "w", encoding="utf-8")as f:
        for i in text_after_sorted:
            f.write(i)


